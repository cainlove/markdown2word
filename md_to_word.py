#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Markdown转Word文档转换工具

这个脚本用于将Markdown格式的文件转换为Word文档格式。
支持基本的Markdown元素，如标题、列表、表格、粗体、斜体等。

使用方法：
    python md_to_word.py <输入的Markdown文件路径> [输出的Word文件路径]

    如果不指定输出路径，将使用输入文件名，但扩展名改为.docx
"""

import sys
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement


def create_element(name):
    """创建XML元素"""
    return OxmlElement(name)


def create_attribute(element, name, value):
    """创建XML属性"""
    element.set(qn(name), value)


def add_page_break(document):
    """添加分页符"""
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def create_table(document, markdown_table):
    """从Markdown表格创建Word表格"""
    # 分割表格行
    rows = [row.strip() for row in markdown_table.strip().split('\n')]
    
    # 处理表头分隔行
    if len(rows) > 1 and re.match(r'^[\|\s\-:]+$', rows[1]):
        header_row = rows[0]
        rows = [header_row] + rows[2:]
    
    # 计算列数
    max_cols = 0
    for row in rows:
        cells = [cell.strip() for cell in row.split('|')]
        # 移除空单元格（表格开头和结尾的|会产生空字符串）
        cells = [cell for cell in cells if cell]
        max_cols = max(max_cols, len(cells))
    
    # 创建表格
    table = document.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    
    # 填充表格内容
    for i, row in enumerate(rows):
        cells = [cell.strip() for cell in row.split('|')]
        # 移除空单元格
        cells = [cell for cell in cells if cell]
        
        for j, cell_content in enumerate(cells):
            if j < max_cols:  # 确保不超出列数
                # 处理单元格中的Markdown格式
                cell = table.cell(i, j)
                paragraph = cell.paragraphs[0]
                process_markdown_text(paragraph, cell_content)
    
    return table


def process_markdown_text(paragraph, text):
    """处理Markdown文本中的格式，如粗体等"""
    # 处理粗体 (**text** 或 __text__)
    # 注意：标题中的双星号会在标题处理部分单独处理，这里处理的是普通文本中的粗体
    bold_pattern = r'(\*\*|__)(.+?)\1'
    
    # 去除文本中的双星号，保留内容
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    
    # 处理行内代码 (`text`)
    code_pattern = r'`(.+?)`'
    
    # 处理链接 [text](url)
    link_pattern = r'\[(.+?)\]\((.+?)\)'
    
    # 使用正则表达式查找所有格式标记
    formats = []
    
    # 查找粗体
    for match in re.finditer(bold_pattern, text):
        formats.append({
            'start': match.start(),
            'end': match.end(),
            'content': match.group(2),
            'type': 'bold',
            'original': match.group(0)
        })
    
    # 不处理斜体，按照要求所有文本不倾斜
    
    # 查找行内代码
    for match in re.finditer(code_pattern, text):
        # 确保不与其他格式重叠
        is_valid = True
        for fmt in formats:
            if (match.start() >= fmt['start'] and match.start() < fmt['end']) or \
               (match.end() > fmt['start'] and match.end() <= fmt['end']):
                is_valid = False
                break
        
        if is_valid:
            formats.append({
                'start': match.start(),
                'end': match.end(),
                'content': match.group(1),
                'type': 'code',
                'original': match.group(0)
            })
    
    # 查找链接
    for match in re.finditer(link_pattern, text):
        # 确保不与其他格式重叠
        is_valid = True
        for fmt in formats:
            if (match.start() >= fmt['start'] and match.start() < fmt['end']) or \
               (match.end() > fmt['start'] and match.end() <= fmt['end']):
                is_valid = False
                break
        
        if is_valid:
            formats.append({
                'start': match.start(),
                'end': match.end(),
                'content': match.group(1),
                'url': match.group(2),
                'type': 'link',
                'original': match.group(0)
            })
    
    # 按起始位置排序
    formats.sort(key=lambda x: x['start'])
    
    # 如果没有格式标记，直接添加文本
    if not formats:
        run = paragraph.add_run(text)
        run.font.name = '宋体'
        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        return
    
    # 处理文本和格式
    last_end = 0
    for fmt in formats:
        # 添加格式标记前的普通文本
        if fmt['start'] > last_end:
            run = paragraph.add_run(text[last_end:fmt['start']])
            run.font.name = '宋体'
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        
        # 添加带格式的文本
        run = paragraph.add_run(fmt['content'])
        
        # 所有文本统一为宋体、黑色
        run.font.name = '宋体'
        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        
        if fmt['type'] == 'bold':
            run.bold = True
        elif fmt['type'] == 'code':
            # 代码块保持等宽字体但颜色改为黑色
            run.font.name = 'Courier New'
        elif fmt['type'] == 'link':
            # 添加超链接，但颜色改为黑色
            run.underline = True
            # 注意：python-docx对超链接的支持有限，这里只是模拟超链接的外观
        
        last_end = fmt['end']
    
    # 添加最后一个格式标记后的普通文本
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = '宋体'
        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色


def convert_markdown_to_word(md_file_path, docx_file_path=None):
    """将Markdown文件转换为Word文档"""
    # 如果未指定输出文件路径，使用输入文件名但扩展名改为.docx
    if not docx_file_path:
        docx_file_path = os.path.splitext(md_file_path)[0] + '.docx'
    
    # 创建一个新的Word文档
    document = Document()
    
    # 设置默认字体为宋体、4号字（约14磅）
    style = document.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(14)  # 4号字约为14磅
    
    # 读取Markdown文件
    with open(md_file_path, 'r', encoding='utf-8') as file:
        md_content = file.read()
    
    # 分割Markdown内容为行
    lines = md_content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # 处理标题 (# 标题)
        if line.startswith('#'):
            # 计算标题级别 (# = 1, ## = 2, ...)
            level = 0
            for char in line:
                if char == '#':
                    level += 1
                else:
                    break
            
            if level <= 6:  # Word支持6级标题
                # 提取标题文本
                title_text = line[level:].strip()
                
                # 去除标题文本中的双星号（**）
                title_text = re.sub(r'\*\*(.+?)\*\*', r'\1', title_text)
                
                # 添加标题，设置为黑体加粗、二号字（约29磅）
                heading = document.add_heading(level=level)
                run = heading.add_run(title_text)
                run.font.name = '黑体'
                run.bold = True
                run.font.size = Pt(29)  # 二号字约为29磅
                run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            else:
                # 如果#超过6个，作为普通段落处理
                paragraph = document.add_paragraph()
                process_markdown_text(paragraph, line)
        
        # 处理水平线 (---, ***, ___)
        elif re.match(r'^(---|\*\*\*|___)$', line):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.border_bottom = True
            paragraph.paragraph_format.space_after = Pt(12)
        
        # 处理表格
        elif line.startswith('|') and i + 1 < len(lines) and lines[i+1].strip().startswith('|'):
            # 收集表格的所有行
            table_lines = [line]
            j = i + 1
            
            while j < len(lines) and lines[j].strip().startswith('|'):
                table_lines.append(lines[j].strip())
                j += 1
            
            # 创建表格
            table_md = '\n'.join(table_lines)
            create_table(document, table_md)
            
            # 更新索引
            i = j - 1
        
        # 处理无序列表 (*, -, +)
        elif re.match(r'^[\*\-\+]\s', line):
            # 收集列表的所有项
            list_items = []
            current_level = 0
            j = i
            
            while j < len(lines) and re.match(r'^\s*[\*\-\+]\s', lines[j]):
                # 计算缩进级别
                indent = len(lines[j]) - len(lines[j].lstrip())
                list_items.append((lines[j].strip()[2:], indent))
                j += 1
            
            # 添加列表项
            for item_text, indent in list_items:
                paragraph = document.add_paragraph(style='List Bullet')
                process_markdown_text(paragraph, item_text)
                # 注意：python-docx对多级列表的支持有限
            
            # 更新索引
            i = j - 1
        
        # 处理有序列表 (1. 2. 3. ...)
        elif re.match(r'^\d+\.\s', line):
            # 收集列表的所有项
            list_items = []
            j = i
            
            while j < len(lines) and re.match(r'^\s*\d+\.\s', lines[j]):
                list_items.append(lines[j].strip()[3:])
                j += 1
            
            # 添加列表项
            for item_text in list_items:
                paragraph = document.add_paragraph(style='List Number')
                process_markdown_text(paragraph, item_text)
            
            # 更新索引
            i = j - 1
        
        # 处理代码块
        elif line.startswith('```'):
            # 收集代码块的所有行
            code_lines = []
            j = i + 1
            
            while j < len(lines) and not lines[j].startswith('```'):
                code_lines.append(lines[j])
                j += 1
            
            # 添加代码块
            code_text = '\n'.join(code_lines)
            paragraph = document.add_paragraph()
            run = paragraph.add_run(code_text)
            run.font.name = '宋体'
            run.font.size = Pt(14)  # 4号字
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            paragraph.paragraph_format.keep_together = True
            
            # 更新索引
            i = j
        
        # 处理引用块 (> ...)
        elif line.startswith('>'):
            # 收集引用块的所有行
            quote_lines = []
            j = i
            
            while j < len(lines) and lines[j].strip().startswith('>'):
                quote_text = lines[j].strip()[1:].strip()
                quote_lines.append(quote_text)
                j += 1
            
            # 添加引用块
            quote_text = ' '.join(quote_lines)
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.right_indent = Inches(0.5)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            process_markdown_text(paragraph, quote_text)
            
            # 更新索引
            i = j - 1
        
        # 处理普通段落
        elif line:
            paragraph = document.add_paragraph()
            process_markdown_text(paragraph, line)
        
        # 处理空行
        else:
            document.add_paragraph()
        
        i += 1
    
    # 保存Word文档
    document.save(docx_file_path)
    return docx_file_path


def main():
    """主函数"""
    # 检查命令行参数
    if len(sys.argv) < 2:
        print(f"使用方法: {sys.argv[0]} <输入的Markdown文件路径> [输出的Word文件路径]")
        sys.exit(1)
    
    # 获取输入文件路径
    md_file_path = sys.argv[1]
    
    # 检查输入文件是否存在
    if not os.path.exists(md_file_path):
        print(f"错误: 文件 '{md_file_path}' 不存在")
        sys.exit(1)
    
    # 获取输出文件路径（如果提供）
    docx_file_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        # 转换Markdown到Word
        output_path = convert_markdown_to_word(md_file_path, docx_file_path)
        print(f"转换成功! Word文档已保存到: {output_path}")
    except Exception as e:
        print(f"转换过程中出错: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()